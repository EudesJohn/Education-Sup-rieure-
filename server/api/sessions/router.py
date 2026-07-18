"""Routeur pour la gestion des sessions d'examen."""

import hashlib
import json
import logging
import random
import string
import traceback
from pathlib import Path
from uuid import uuid4

import io

from fastapi import APIRouter, Body, Depends, HTTPException, Query, UploadFile, File, Form, status
from fastapi.responses import StreamingResponse

from core.db import (
    get_session_by_id,
    create_session,
    update_session,
    delete_session,
    get_teacher_sessions,
    get_session_exams,
    get_variants_by_exercise,
    create_generated_exam,
    create_exercise,
    create_variant,
    get_session_exercises,
    add_session_exercise,
    remove_session_exercise,
    update_session_exercise_order,
    list_class_students,
    get_list_entries,
)
from core.security import hash_student_identifier
from core.dependencies import get_current_teacher
from core.supabase_client import get_supabase
from schemas.sessions import (
    ExamSessionCreate,
    ExamSessionUpdate,
    ExamSessionResponse,
    ExamGenerateRequest,
    SessionExerciseAdd,
    SessionExerciseReorder,
    SessionExerciseResponse,
)
from services.qcm_generator import QCMGenerator
from services.content_parser import parse_exam_content, detect_language

logger = logging.getLogger(__name__)

router = APIRouter()


def _redistribute_points(questions: list[dict], total_score: float) -> list[dict]:
    """Redistribue les points uniformement entre les questions.

    Garantit que la somme des points = total_score exactement (tolerance 0.00).
    Le dernier exercice recoit le reste pour compenser les arrondis.
    """
    if not questions:
        return questions
    n = len(questions)
    base = round(total_score / n, 2)
    somme = round(base * n, 2)
    diff = round(total_score - somme, 2)
    for i, q in enumerate(questions):
        if i == n - 1:
            q["points"] = round(base + diff, 2)
        else:
            q["points"] = base
    return questions


def _generate_access_code() -> str:
    """Génère un code d'accès aléatoire unique pour une session."""
    chars = string.ascii_uppercase + string.digits
    # 8 caractères â†’ 36^8 = 2 821 109 907 456 combinaisons
    return "".join(random.choices(chars, k=8))


@router.get("")
def list_sessions(
    skip: int = Query(0, ge=0),
    limit: int = Query(50, ge=1, le=200),
    teacher: dict = Depends(get_current_teacher),
):
    """Lister les sessions de l'enseignant (pagine)."""
    sessions = get_teacher_sessions(teacher["id"])
    total = len(sessions)
    paginated = sessions[skip : skip + limit]

    return {
        "items": [ExamSessionResponse.model_validate(s) for s in paginated],
        "total": total,
        "skip": skip,
        "limit": limit,
    }


@router.post("", response_model=ExamSessionResponse, status_code=201)
def create_session_route(
    data: ExamSessionCreate,
    teacher: dict = Depends(get_current_teacher),
):
    """Creer une nouvelle session d'examen."""
    session_data = data.model_dump()
    session_data["teacher_id"] = teacher["id"]

    # Generer un code d'acces unique
    supabase = get_supabase()
    for _attempt in range(10):
        code = _generate_access_code()
        existing = supabase.table("exam_sessions").select("id").eq("access_code", code).maybe_single().execute()
        if not existing or not existing.data:
            session_data["access_code"] = code
            break
    else:
        raise HTTPException(status_code=500, detail="Erreur lors de la génération du code d'accès")

    created = create_session(session_data)
    if not created:
        raise HTTPException(status_code=500, detail="Erreur lors de la création de la session")
    return ExamSessionResponse.model_validate(created)


@router.post("/{session_id}/generate-exams", status_code=status.HTTP_201_CREATED)
def generate_exams(
    session_id: int,
    data: ExamGenerateRequest,
    teacher: dict = Depends(get_current_teacher),
):
    """Genere des epreuves uniques pour chaque etudiant a partir des exercices selectionnes.

    Pour chaque etudiant, le moteur tire aleatoirement une combinaison de variantes,
    garantissant que chaque epreuve est unique.

    Si exercise_ids est omis, utilise les exercices de la table session_exercises.
    """
    session = get_session_by_id(session_id)
    if not session or session["teacher_id"] != teacher["id"]:
        raise HTTPException(status_code=404, detail="Session non trouvée")
    if session["status"] != "draft":
        raise HTTPException(
            status_code=400,
            detail="Seules les sessions en brouillon peuvent recevoir de nouvelles épreuves",
        )

    # Si exercise_ids non fourni, utiliser session_exercises
    exercise_ids = data.exercise_ids
    if not exercise_ids:
        linked = get_session_exercises(session_id)
        if not linked:
            raise HTTPException(
                status_code=400,
                detail="Aucun exercice lié à cette session. "
                "Ajoutez des exercices via POST /sessions/{id}/exercises "
                "ou fournissez exercise_ids dans le body.",
            )
        exercise_ids = [l["exercise_id"] for l in linked]
        # Charger les points_override s'ils existent
        points_overrides = {l["exercise_id"]: l["points_override"] for l in linked if l.get("points_override") is not None}
    else:
        points_overrides = {}

    # Recuperer les exercices avec leurs variantes
    supabase = get_supabase()
    exercises_result = (
        supabase.table("exercises")
        .select("*")
        .in_("id", exercise_ids)
        .eq("teacher_id", teacher["id"])
        .execute()
    )
    exercises = exercises_result.data

    if len(exercises) != len(exercise_ids):
        found_ids = {ex["id"] for ex in exercises}
        missing = set(exercise_ids) - found_ids
        raise HTTPException(
            status_code=404,
            detail=f"Exercices introuvables : {missing}",
        )

    # Verifier que tous les exercices ont des variantes et les charger
    # Si un exercice n'a pas de variante, en créer une par défaut automatiquement
    for ex in exercises:
        variants = get_variants_by_exercise(ex["id"])
        if not variants:
            # Creer une variante par défaut
            default_content = ex.get("content") or ex.get("instructions") or "Ã‰noncé par défaut"
            create_variant({
                "exercise_id": ex["id"],
                "variant_order": 0,
                "content": default_content,
                "data_overrides": None,
            })
            variants = get_variants_by_exercise(ex["id"])
            logger.warning("Variante par défaut créée pour l'exercice '%s' (id=%s)", ex["title"], ex["id"])
        ex["_variants"] = variants

    # Preparer les identifiants etudiants
    if data.student_identifiers:
        student_ids = data.student_identifiers
        if len(student_ids) != session["student_count"]:
            raise HTTPException(
                status_code=400,
                detail=f"Nombre d'etudiants fourni ({len(student_ids)}) different du nombre declare ({session['student_count']})",
            )
    else:
        # Chercher les vrais étudiants depuis class_id ou student_list_id
        real_students = []
        if session.get("class_id"):
            real_students = list_class_students(session["class_id"])
        elif session.get("student_list_id"):
            real_students = get_list_entries(session["student_list_id"])

        if real_students:
            student_ids = [
                {
                    "student_name": s["student_name"],
                    "student_number": s["student_number"],
                }
                for s in real_students
            ]
        else:
            # Fallback : auto-generer des identifiants sequentiels
            student_ids = [
                {
                    "student_name": f"Etudiant {i + 1}",
                    "student_number": f"PEAN_{session['id']}_{i + 1:04d}",
                }
                for i in range(session["student_count"])
            ]

    # Generer les epreuves uniques
    used_combinations: set[str] = set()
    generated_exams = []

    for student_info in student_ids:
        # Mélanger l'ordre des exercices pour cet étudiant
        student_exercises = list(exercises)
        random.shuffle(student_exercises)

        # Tirer aléatoirement une combinaison de variantes
        assignment: dict[int, dict] = {}
        chosen_assignment = None

        for _attempt in range(50):
            temp_assignment: dict[int, dict] = {}
            combo_parts = []
            for ex in student_exercises:
                variants = ex.get("_variants", [])
                if not variants:
                    variant = {
                        "id": 0, "variant_order": 0,
                        "content": ex.get("instructions", ""),
                        "data_overrides": None,
                    }
                else:
                    variant = random.choice(variants)
                temp_assignment[ex["id"]] = variant
                # combo_parts contient l'ordre de la question et l'ID de sa variante
                combo_parts.append(f"{ex['id']}:{variant['id']}")

            combo_key = ":".join(combo_parts)
            if combo_key not in used_combinations:
                used_combinations.add(combo_key)
                chosen_assignment = temp_assignment
                break

        if chosen_assignment is not None:
            assignment = chosen_assignment
        else:
            # Fallback si pas de combinaison unique trouvée après 50 tentatives (recyclage)
            assignment = {}
            for ex in student_exercises:
                variants = ex.get("_variants", [])
                if not variants:
                    assignment[ex["id"]] = {
                        "id": 0, "variant_order": 0,
                        "content": ex.get("instructions", ""),
                        "data_overrides": None,
                    }
                else:
                    assignment[ex["id"]] = random.choice(variants)

        # Assembler le contenu JSON de l'epreuve
        content_parts = []
        for ex in student_exercises:
            variant = assignment[ex["id"]]
            ex_points = points_overrides.get(ex["id"]) if ex["id"] in points_overrides else ex["points"]
            content_parts.append({
                "exercise_id": ex["id"],
                "exercise_title": ex["title"],
                "difficulty": ex["difficulty"],
                "points": ex_points,
                "instructions": ex["instructions"],
                "exercise_type": ex["exercise_type"],
                "language": ex.get("language"),
                "variant_id": variant["id"],
                "variant_order": variant["variant_order"],
                "content": variant["content"],
                "data_overrides": json.loads(variant.get("data_overrides") or "null"),
            })

        content_json = json.dumps(content_parts, ensure_ascii=False)

        # Hashes
        student_hash = hash_student_identifier(session["id"], student_info["student_number"])

        variant_ids = sorted(v["id"] for v in assignment.values())
        combo_raw = f"{session['id']}:{variant_ids}"
        variant_combo_hash = hashlib.sha256(combo_raw.encode()).hexdigest()
        # sha256_hash doit être unique par étudiant â€” inclure le hash étudiant
        sha256_hash = hashlib.sha256(f"{combo_raw}:{student_hash}".encode()).hexdigest()


        exam_data = {
            "session_id": session["id"],
            "student_id_hash": student_hash,
            "variant_combo_hash": variant_combo_hash,
            "sha256_hash": sha256_hash,
            "content": content_json,
            "status": "pending",
        }
        created_exam = create_generated_exam(exam_data)
        if created_exam:
            generated_exams.append(created_exam)

    # Calculer le nombre maximum de combinaisons (comprenant les variantes et l'ordre des questions)
    import math
    max_combinations = math.factorial(len(exercises))
    for ex in exercises:
        variants_count = len(ex.get("_variants", []))
        if variants_count > 0:
            max_combinations *= variants_count

    return {
        "generated": len(generated_exams),
        "message": f"{len(generated_exams)} epreuves uniques generees avec succes",
        "max_combinations": max_combinations,
        "exams": [
            {
                "id": e["id"],
                "sha256_hash": e["sha256_hash"],
                "status": e["status"],
            }
            for e in generated_exams
        ],
    }


@router.post("/{session_id}/generate-qcm-ai")
async def generate_qcm_ai(
    session_id: int,
    teacher: dict = Depends(get_current_teacher),
    file: UploadFile = File(None),
    text_content: str = Form(None),
    num_questions: int = Form(5),
    exercise_type: str = Form("mixed"),
):
    """Generer des exercices par IA a partir d'un fichier (PDF/Word) ou d'un texte saisi.

    L'IA analyse le contenu et produit des questions avec variantes,
    les enregistre dans la banque d'exercices, puis genere les epreuves uniques.
    """
    session = get_session_by_id(session_id)
    if not session or session["teacher_id"] != teacher["id"]:
        raise HTTPException(status_code=404, detail="Session non trouvée")
    if session["status"] != "draft":
        raise HTTPException(status_code=400, detail="Seules les sessions en brouillon peuvent recevoir des exercices")

    # Extraire le contenu texte
    content = None
    source_label = "texte saisi"

    if file:
        source_label = file.filename or "fichier"
        raw = await file.read()
        ext = file.filename.split(".")[-1].lower() if file.filename else ""

        if ext in ("txt", "md", "html", "htm"):
            content = raw.decode("utf-8", errors="replace")
        elif ext in ("pdf",):
            # Extraction PDF simple via PyMuPDF si disponible
            try:
                import fitz
                doc = fitz.open(stream=raw, filetype="pdf")
                content = "\n".join(page.get_text() for page in doc)
            except ImportError:
                raise HTTPException(status_code=400, detail="PyMuPDF (fitz) requis pour l'extraction PDF")
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Erreur d'extraction PDF: {e}")
        elif ext in ("docx", "doc"):
            try:
                import docx
                doc = docx.Document(raw)
                content = "\n".join(p.text for p in doc.paragraphs)
            except ImportError:
                raise HTTPException(status_code=400, detail="python-docx requis pour l'extraction Word")
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Erreur d'extraction Word: {e}")
        else:
            # Fallback: essayer de decoder en texte
            try:
                content = raw.decode("utf-8", errors="replace")
            except Exception:
                raise HTTPException(status_code=400, detail=f"Format de fichier non supporte: .{ext}")
    elif text_content and text_content.strip():
        content = text_content.strip()
    else:
        raise HTTPException(status_code=400, detail="Fournissez un fichier ou un texte")

    if not content or len(content.strip()) < 20:
        raise HTTPException(status_code=400, detail="Contenu trop court (min 20 caracteres)")

    # Generer les exercices via IA
    generator = QCMGenerator()
    grading_system = session.get("grading_system", "20")
    total_score = int(grading_system) if grading_system.isdigit() else 20
    result = await generator.generate(
        content,
        num_questions=num_questions,
        exercise_type=exercise_type,
        total_score=total_score,
    )

    if "error" in result:
        raise HTTPException(status_code=502, detail=result["error"])

    questions = result["questions"]
    # Redistribuer les points pour garantir somme = total_score (tolerance 0.00)
    questions = _redistribute_points(questions, total_score)
    warnings = generator.validate_questions(questions)

    # Enregistrer chaque question comme exercice + variantes
    created_exercises = []
    for q in questions:
        exercise_data = {
            "teacher_id": teacher["id"],
            "title": q.get("title", f"Question {len(created_exercises) + 1}"),
            "subject": q.get("subject", session["subject"]),
            "difficulty": q.get("difficulty", "medium"),
            "instructions": q.get("instructions", ""),
            "correct_answer": q.get("correct_answer", ""),
            "points": q.get("points", 10),
            "exercise_type": q.get("exercise_type", exercise_type),
        }
        ex = create_exercise(exercise_data)
        if not ex:
            continue

        # Creer les variantes
        for v in q.get("variants", []):
            variant_data = {
                "exercise_id": ex["id"],
                "variant_order": v.get("variant_order", 0),
                "content": v.get("content", ""),
                "data_overrides": json.dumps(v.get("data_overrides")) if v.get("data_overrides") else None,
            }
            create_variant(variant_data)

        created_exercises.append({
            "id": ex["id"],
            "title": ex["title"],
            "variants_count": len(q.get("variants", [])),
        })

    if not created_exercises:
        raise HTTPException(status_code=502, detail="Aucun exercice n'a pu etre cree depuis la reponse IA")

    return {
        "generated": len(created_exercises),
        "exercises": created_exercises,
        "warnings": warnings,
        "source": source_label,
        "message": f"{len(created_exercises)} questions QCM generees depuis '{source_label}'",
    }


@router.post("/{session_id}/upload-exam", status_code=201)
async def upload_exam_file(
    session_id: int,
    teacher: dict = Depends(get_current_teacher),
    file: UploadFile = File(None),
    text_content: str = Form(None),
    num_questions: int = Form(3),  # max 3 pour rester dans le timeout Vercel
    exercise_type: str = Form("mixed"),
):
    """Uploader un sujet â†’ IA lit le contenu â†’ genere les epreuves individuelles pour chaque etudiant.

    Le professeur fournit un fichier (PDF/Word/TXT) ou un texte.
    L'IA extrait le contenu, cree des exercices avec variantes uniques,
    puis genere des epreuves personnalisees pour chaque etudiant.
    """
    try:
        session = get_session_by_id(session_id)
        if not session or session["teacher_id"] != teacher["id"]:
            raise HTTPException(status_code=404, detail="Session non trouvée")
        if session["status"] != "draft":
            raise HTTPException(
                status_code=400,
                detail="Seules les sessions en brouillon peuvent recevoir des exercices",
            )

        # 1. Extraire le contenu texte
        content = None
        source_label = "texte saisi"

        if file:
            source_label = file.filename or "fichier"
            raw = await file.read()
            ext = file.filename.split(".")[-1].lower() if file.filename else ""

            if ext in ("txt", "md", "html", "htm"):
                content = raw.decode("utf-8", errors="replace")
            elif ext in ("pdf",):
                try:
                    import fitz
                    doc = fitz.open(stream=raw, filetype="pdf")
                    content = "\n".join(page.get_text() for page in doc)
                except ImportError:
                    raise HTTPException(status_code=400, detail="PyMuPDF (fitz) requis pour l'extraction PDF")
                except Exception as e:
                    raise HTTPException(status_code=400, detail=f"Erreur d'extraction PDF: {e}")
            elif ext in ("docx", "doc"):
                try:
                    import docx
                    doc = docx.Document(raw)
                    content = "\n".join(p.text for p in doc.paragraphs)
                except ImportError:
                    raise HTTPException(status_code=400, detail="python-docx requis pour l'extraction Word")
                except Exception as e:
                    raise HTTPException(status_code=400, detail=f"Erreur d'extraction Word: {e}")
            else:
                try:
                    content = raw.decode("utf-8", errors="replace")
                except Exception:
                    raise HTTPException(status_code=400, detail=f"Format de fichier non supporte: .{ext}")
        elif text_content and text_content.strip():
            content = text_content.strip()
        else:
            raise HTTPException(status_code=400, detail="Fournissez un fichier ou un texte")

        if not content or len(content.strip()) < 20:
            raise HTTPException(status_code=400, detail="Contenu trop court (min 20 caracteres)")

        # 2. Generer les exercices via IA
        generator = QCMGenerator()
        grading_system = session.get("grading_system", "20")
        total_score = int(grading_system) if grading_system.isdigit() else 20
        result = await generator.generate(
            content,
            num_questions=num_questions,
            exercise_type=exercise_type,
            total_score=total_score,
        )

        if "error" in result:
            raise HTTPException(status_code=502, detail=result["error"])

        questions = result["questions"]
        # Redistribuer les points pour garantir somme = total_score (tolerance 0.00)
        questions = _redistribute_points(questions, total_score)
        warnings = result.get("warnings", [])

        # 3. Creer les exercices + variantes et les lier a la session
        supabase = get_supabase()
        created_exercises = []
        sort_order = 0

        for q in questions:
            exercise_data = {
                "teacher_id": teacher["id"],
                "title": q.get("title", f"Question {sort_order + 1}"),
                "subject": q.get("subject", session["subject"]),
                "difficulty": q.get("difficulty", "medium"),
                "instructions": q.get("instructions", ""),
                "correct_answer": q.get("correct_answer", ""),
                "points": q.get("points", 10),
                "exercise_type": q.get("exercise_type", exercise_type),
            }
            ex = create_exercise(exercise_data)
            if not ex:
                continue

            # Creer les variantes
            for v in q.get("variants", []):
                variant_data = {
                    "exercise_id": ex["id"],
                    "variant_order": v.get("variant_order", 0),
                    "content": v.get("content", ""),
                    "data_overrides": json.dumps(v.get("data_overrides")) if v.get("data_overrides") else None,
                }
                create_variant(variant_data)

            # Lier l'exercice a la session
            add_session_exercise(
                session_id=session_id,
                exercise_id=ex["id"],
                sort_order=sort_order,
                points_override=q.get("points"),
            )
            sort_order += 1

            created_exercises.append({
                "id": ex["id"],
                "title": ex["title"],
                "variants_count": len(q.get("variants", [])),
            })

        if not created_exercises:
            raise HTTPException(status_code=502, detail="Aucun exercice n'a pu etre cree depuis la reponse IA")

        # 4. Supprimer les anciennes epreuves generees
        existing_exams = get_session_exams(session_id)
        for exam in existing_exams:
            supabase.table("generated_exams").delete().eq("id", exam["id"]).execute()

        # 5. Generer les epreuves individuelles pour chaque etudiant
        # Recuperer les exercices lies avec leurs variantes
        linked = get_session_exercises(session_id)
        exercise_ids = [l["exercise_id"] for l in linked]
        points_overrides = {l["exercise_id"]: l["points_override"] for l in linked if l.get("points_override") is not None}

        exercises_result = (
            supabase.table("exercises")
            .select("*")
            .in_("id", exercise_ids)
            .execute()
        )
        all_exercises = exercises_result.data

        if len(all_exercises) != len(exercise_ids):
            raise HTTPException(status_code=404, detail="Certains exercices ne sont plus disponibles")

        # Charger les variantes (créer une variante par défaut si aucune n'existe)
        for ex in all_exercises:
            variants = get_variants_by_exercise(ex["id"])
            if not variants:
                default_content = ex.get("content") or ex.get("instructions") or "Ã‰noncé par défaut"
                create_variant({
                    "exercise_id": ex["id"],
                    "variant_order": 0,
                    "content": default_content,
                    "data_overrides": None,
                })
                variants = get_variants_by_exercise(ex["id"])
                logger.warning("Variante par défaut créée pour l'exercice '%s' (id=%s)", ex["title"], ex["id"])
            ex["_variants"] = variants

        # Recuperer les etudiants
        student_ids = []
        if session.get("class_id"):
            real_students = list_class_students(session["class_id"])
            student_ids = [{"student_name": s["student_name"], "student_number": s["student_number"]} for s in real_students]
        elif session.get("student_list_id"):
            entries = get_list_entries(session["student_list_id"])
            student_ids = [{"student_name": e["student_name"], "student_number": e["student_number"]} for e in entries]
        else:
            count = session.get("student_count") or 0
            student_ids = [{"student_name": f"Etudiant {i + 1}", "student_number": f"PEAN_{session['id']}_{i + 1:04d}"} for i in range(count)]

        if not student_ids:
            raise HTTPException(
                status_code=400,
                detail="Aucun étudiant dans cette session. Ajoutez des étudiants ou une classe avant de générer les épreuves.",
            )

        import math

        # Separer code / non-code
        code_exercises = [ex for ex in all_exercises if ex.get("exercise_type") == "code"]
        variant_exercises = [ex for ex in all_exercises if ex.get("exercise_type") != "code"]

        code_assignment: dict[int, dict] = {}
        for ex in code_exercises:
            base_variant = ex["_variants"][0] if ex["_variants"] else {"id": 0, "variant_order": 0, "content": ex.get("instructions", ""), "data_overrides": None}
            code_assignment[ex["id"]] = base_variant

        # Calcul du nombre max de combinaisons uniques (variantes x ordre de shuffle)
        n_total = len(all_exercises)
        variant_combos = 1
        for ex in variant_exercises:
            variant_combos *= max(1, len(ex.get("_variants", [])))
        shuffle_combos = math.factorial(n_total) if n_total > 0 else 1
        max_combinations = variant_combos * shuffle_combos

        if max_combinations < len(student_ids):
            raise HTTPException(
                status_code=400,
                detail=f"Pas assez de combinaisons uniques pour {len(student_ids)} etudiants "
                       f"(seulement {max_combinations} possibles). "
                       f"Ajoute des variantes ou reduis le nombre d'etudiants."
            )

        # Distribution avec verification d'unicite
        total_exams_generated = 0
        exam_batch = []
        used_combinations: set[str] = set()

        for student_info in student_ids:
            for _attempt in range(100):
                student_exercises = list(all_exercises)
                random.shuffle(student_exercises)

                assignment: dict[int, dict] = {}
                combo_parts: list[str] = []

                for ex in student_exercises:
                    if ex.get("exercise_type") == "code":
                        variant = code_assignment[ex["id"]]
                    else:
                        variants = ex.get("_variants", [])
                        variant = random.choice(variants) if variants else {"id": 0, "variant_order": 0, "content": ex.get("instructions", ""), "data_overrides": None}
                    assignment[ex["id"]] = variant
                    combo_parts.append(f"{ex['id']}:{variant['id']}")

                combo_key = ":".join(combo_parts)
                if combo_key not in used_combinations:
                    used_combinations.add(combo_key)
                    break
            else:
                raise HTTPException(
                    status_code=400,
                    detail="Impossible de trouver une combinaison unique pour chaque etudiant. "
                           "Ajoute des variantes ou reduis le nombre d'etudiants."
                )

            # Construire le contenu dans l'ordre shuffle de cet etudiant
            content_parts = []
            for ex in student_exercises:
                ex_points = points_overrides.get(ex["id"], ex["points"])
                variant = assignment[ex["id"]]
                content_parts.append({
                    "exercise_id": ex["id"],
                    "exercise_title": ex["title"],
                    "difficulty": ex["difficulty"],
                    "points": ex_points,
                    "instructions": ex.get("instructions", ""),
                    "exercise_type": ex.get("exercise_type", "qcm"),
                    "language": ex.get("language"),
                    "variant_id": variant["id"],
                    "variant_order": variant["variant_order"],
                    "content": variant["content"],
                    "data_overrides": json.loads(variant.get("data_overrides") or "null"),
                })

            content_json = json.dumps(content_parts, ensure_ascii=False)
            student_hash = hash_student_identifier(session["id"], student_info["student_number"])
            variant_ids_sorted = sorted(v["id"] for v in assignment.values())
            combo_raw = f"{session['id']}:{variant_ids_sorted}"
            variant_combo_hash = hashlib.sha256(combo_raw.encode()).hexdigest()
            sha256_hash = hashlib.sha256(f"{combo_raw}:{student_hash}".encode()).hexdigest()

            exam_batch.append({
                "session_id": session["id"],
                "student_id_hash": student_hash,
                "variant_combo_hash": variant_combo_hash,
                "sha256_hash": sha256_hash,
                "content": content_json,
                "status": "pending",
            })

            if len(exam_batch) >= 50:
                supabase.table("generated_exams").insert(exam_batch).execute()
                total_exams_generated += len(exam_batch)
                exam_batch = []

        if exam_batch:
            supabase.table("generated_exams").insert(exam_batch).execute()
            total_exams_generated += len(exam_batch)

        return {
            "generated": total_exams_generated,
            "exercises_created": len(created_exercises),
            "warnings": warnings,
            "source": source_label,
            "message": f"{total_exams_generated} epreuves generees depuis '{source_label}'",
        }
    except HTTPException:
        raise
    except Exception as e:
        tb = traceback.format_exc()
        logger.critical("ERREUR upload_exam_file session=%s: %s\n%s", session_id, e, tb)
        raise HTTPException(status_code=500, detail=f"Erreur interne: {type(e).__name__}: {e}")


@router.post("/{session_id}/upload-exam-json", status_code=201)
async def upload_exam_json(
    session_id: int,
    teacher: dict = Depends(get_current_teacher),
    data: dict = Body(...),
):
    """Version JSON de upload-exam - contourne les problemes FormData/502 du navigateur.

    Corps JSON : {"text_content": "...", "num_questions": 3, "exercise_type": "qcm"}
    """
    try:
        return await _do_upload_exam_json(session_id, teacher, data)
    except HTTPException:
        raise
    except Exception as exc:
        tb = traceback.format_exc()
        logger.exception("Erreur inattendue dans upload_exam_json (session=%s): %s", session_id, exc)
        raise HTTPException(status_code=500, detail=f"Erreur interne: {type(exc).__name__}: {exc}\n---\n{tb}")


async def _do_upload_exam_json(session_id: int, teacher: dict, data: dict) -> dict:
    """Logique interne de upload-exam-json."""
    if data is None:
        data = {}
    text_content = data.get("text_content") or ""
    num_questions = int(data.get("num_questions", 3))
    exercise_type = data.get("exercise_type", "mixed")

    # Appelle la meme logique que upload_exam_file en reconstituant les parametres
    from core.db import get_session_by_id
    session = get_session_by_id(session_id)
    if not session or session["teacher_id"] != teacher["id"]:
        raise HTTPException(status_code=404, detail="Session non trouvée")
    if session["status"] != "draft":
        raise HTTPException(
            status_code=400,
            detail="Seules les sessions en brouillon peuvent recevoir des exercices",
        )

    if not text_content or len(text_content.strip()) < 20:
        raise HTTPException(status_code=400, detail="Contenu trop court (min 20 caracteres)")

    content = text_content.strip()
    source_label = "texte saisi"

    # 2. Generer les exercices via IA
    generator = QCMGenerator()
    grading_system = session.get("grading_system", "20")
    total_score = int(grading_system) if str(grading_system).isdigit() else 20
    result = await generator.generate(
        content,
        num_questions=num_questions,
        exercise_type=exercise_type,
        total_score=total_score,
    )

    if "error" in result:
        raise HTTPException(status_code=502, detail=result["error"])

    questions = result["questions"]
    # Redistribuer les points pour garantir somme = total_score (tolerance 0.00)
    questions = _redistribute_points(questions, total_score)
    warnings = generator.validate_questions(questions)

    # 3. Creer les exercices + variantes
    supabase = get_supabase()
    created_exercises = []
    sort_order = 0

    for q in questions:
        exercise_data = {
            "teacher_id": teacher["id"],
            "title": q.get("title", f"Question {sort_order + 1}"),
            "subject": q.get("subject", session["subject"]),
            "difficulty": q.get("difficulty", "medium"),
            "instructions": q.get("instructions", ""),
            "correct_answer": q.get("correct_answer", ""),
            "points": q.get("points", 10),
            "exercise_type": q.get("exercise_type", "qcm"),
        }
        ex = create_exercise(exercise_data)
        if not ex:
            continue

        for v in q.get("variants", []):
            variant_data = {
                "exercise_id": ex["id"],
                "variant_order": v.get("variant_order", 0),
                "content": v.get("content", ""),
                "data_overrides": json.dumps(v.get("data_overrides")) if v.get("data_overrides") else None,
            }
            create_variant(variant_data)

        add_session_exercise(
            session_id=session_id,
            exercise_id=ex["id"],
            sort_order=sort_order,
            points_override=q.get("points"),
        )
        sort_order += 1
        created_exercises.append({
            "id": ex["id"],
            "title": ex["title"],
            "variants_count": len(q.get("variants", [])),
        })

    if not created_exercises:
        raise HTTPException(status_code=502, detail="Aucun exercice n'a pu etre cree depuis la reponse IA")

    # 4. Supprimer les anciennes epreuves AVANT d'inserer les nouvelles
    existing_exams = get_session_exams(session_id)
    for exam in existing_exams:
        supabase.table("generated_exams").delete().eq("id", exam["id"]).execute()

    # 5. Generer les epreuves individuelles
    linked = get_session_exercises(session_id)
    exercise_ids = [l["exercise_id"] for l in linked]
    points_overrides = {l["exercise_id"]: l["points_override"] for l in linked if l.get("points_override") is not None}

    exercises_result = (
        supabase.table("exercises")
        .select("*")
        .in_("id", exercise_ids)
        .execute()
    )
    all_exercises = exercises_result.data

    for ex in all_exercises:
        variants = get_variants_by_exercise(ex["id"])
        if not variants:
            default_content = ex.get("content") or ex.get("instructions") or "Ã‰noncé par défaut"
            create_variant({
                "exercise_id": ex["id"],
                "variant_order": 0,
                "content": default_content,
                "data_overrides": None,
            })
            variants = get_variants_by_exercise(ex["id"])
            logger.warning("Variante par défaut créée pour l'exercice '%s' (id=%s)", ex["title"], ex["id"])
        ex["_variants"] = variants

    student_ids = []
    if session.get("class_id"):
        real_students = list_class_students(session["class_id"])
        student_ids = [{"student_name": s["student_name"], "student_number": s["student_number"]} for s in real_students]
    elif session.get("student_list_id"):
        entries = get_list_entries(session["student_list_id"])
        student_ids = [{"student_name": e["student_name"], "student_number": e["student_number"]} for e in entries]
    else:
        count = session.get("student_count") or 0
        student_ids = [{"student_name": f"Etudiant {i + 1}", "student_number": f"PEAN_{session['id']}_{i + 1:04d}"} for i in range(count)]

    if not student_ids:
        raise HTTPException(status_code=400, detail="Aucun etudiant dans cette session")

    import math
    total_exams_generated = 0
    exam_batch = []
    used_combinations: set[str] = set()

    # Calcul du nombre max de combinaisons uniques (variantes x ordre de shuffle)
    n_total = len(all_exercises)
    variant_combos = 1
    for ex in all_exercises:
        variant_combos *= max(1, len(ex.get("_variants", [])))
    shuffle_combos = math.factorial(n_total) if n_total > 0 else 1
    max_combinations = variant_combos * shuffle_combos

    if max_combinations < len(student_ids):
        raise HTTPException(
            status_code=400,
            detail=f"Pas assez de combinaisons uniques pour {len(student_ids)} etudiants "
                   f"(seulement {max_combinations} possibles). "
                   f"Ajoute des variantes ou reduis le nombre d'etudiants."
        )

    for student_info in student_ids:
        for _attempt in range(100):
            student_exercises = list(all_exercises)
            random.shuffle(student_exercises)

            assignment: dict[int, dict] = {}
            combo_parts = []
            for ex in student_exercises:
                variants = ex.get("_variants", [])
                variant = random.choice(variants) if variants else {"id": 0, "variant_order": 0, "content": ex.get("instructions", ""), "data_overrides": None}
                assignment[ex["id"]] = variant
                combo_parts.append(f"{ex['id']}:{variant['id']}")
            combo_key = ":".join(combo_parts)
            if combo_key not in used_combinations:
                used_combinations.add(combo_key)
                break
        else:
            raise HTTPException(
                status_code=400,
                detail="Impossible de trouver une combinaison unique pour chaque etudiant. "
                       "Ajoute des variantes ou reduis le nombre d'etudiants."
            )

        content_parts = []
        for ex in student_exercises:
            ex_points = points_overrides.get(ex["id"], ex["points"])
            variant = assignment[ex["id"]]
            content_parts.append({
                "exercise_id": ex["id"],
                "exercise_title": ex["title"],
                "difficulty": ex["difficulty"],
                "points": ex_points,
                "instructions": ex.get("instructions", ""),
                "exercise_type": ex.get("exercise_type", "qcm"),
                "language": ex.get("language"),
                "variant_id": variant["id"],
                "variant_order": variant.get("variant_order", 0),
                "content": variant["content"],
                "data_overrides": json.loads(variant.get("data_overrides") or "null"),
            })

        content_json = json.dumps(content_parts, ensure_ascii=False)
        student_hash = hash_student_identifier(session["id"], student_info["student_number"])
        variant_ids_sorted = sorted(v["id"] for v in assignment.values())
        combo_raw = f"{session['id']}:{variant_ids_sorted}"
        variant_combo_hash = hashlib.sha256(combo_raw.encode()).hexdigest()
        sha256_hash = hashlib.sha256(f"{combo_raw}:{student_hash}".encode()).hexdigest()

        exam_batch.append({
            "session_id": session["id"],
            "student_id_hash": student_hash,
            "variant_combo_hash": variant_combo_hash,
            "sha256_hash": sha256_hash,
            "content": content_json,
            "status": "pending",
        })

        if len(exam_batch) >= 50:
            supabase.table("generated_exams").insert(exam_batch).execute()
            total_exams_generated += len(exam_batch)
            exam_batch = []

    if exam_batch:
        supabase.table("generated_exams").insert(exam_batch).execute()
        total_exams_generated += len(exam_batch)
    return {
        "generated": total_exams_generated,
        "exercises_created": len(created_exercises),
        "warnings": warnings,
        "source": source_label,
        "message": f"{total_exams_generated} epreuves generees depuis '{source_label}'",
    }


@router.post("/{session_id}/publish-content", status_code=201)
async def publish_shared_content(
    session_id: int,
    teacher: dict = Depends(get_current_teacher),
    file: UploadFile = File(None),
    text_content: str = Form(None),
    exercises_config: str = Form(None),
):
    """Publie un contenu identique pour tous les etudiants (mode partage).

    Pas de generation IA, pas de variantes. Le contenu (fichier ou texte)
    est stocke tel quel dans l'epreuve de chaque etudiant.
    """
    session = get_session_by_id(session_id)
    if not session or session["teacher_id"] != teacher["id"]:
        raise HTTPException(status_code=404, detail="Session non trouvee")
    if session["status"] != "draft":
        raise HTTPException(status_code=400, detail="Seules les sessions en brouillon peuvent recevoir du contenu")

    # 1. Extraire le contenu texte
    content = None
    source_label = "texte saisi"

    if file:
        source_label = file.filename or "fichier"
        raw = await file.read()
        ext = file.filename.split(".")[-1].lower() if file.filename else ""
        if ext in ("txt", "md", "html", "htm"):
            content = raw.decode("utf-8", errors="replace")
        elif ext in ("pdf",):
            try:
                import fitz
                doc = fitz.open(stream=raw, filetype="pdf")
                content = "\n".join(page.get_text() for page in doc)
            except ImportError:
                raise HTTPException(status_code=400, detail="PyMuPDF (fitz) requis pour l'extraction PDF")
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Erreur d'extraction PDF: {e}")
        elif ext in ("docx", "doc"):
            try:
                import docx
                doc = docx.Document(raw)
                content = "\n".join(p.text for p in doc.paragraphs)
            except ImportError:
                raise HTTPException(status_code=400, detail="python-docx requis pour l'extraction Word")
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Erreur d'extraction Word: {e}")
        else:
            content = raw.decode("utf-8", errors="replace")
    elif text_content and text_content.strip():
        content = text_content.strip()
    else:
        raise HTTPException(status_code=400, detail="Fournissez un fichier ou un texte")

    if not content or len(content.strip()) < 20:
        raise HTTPException(status_code=400, detail="Contenu trop court (min 20 caracteres)")

    # 2. Mettre a jour le mode de la session
    update_session(session_id, {"exam_mode": "shared"})

    # 3. Supprimer les anciennes epreuves
    supabase = get_supabase()
    existing_exams = get_session_exams(session_id)
    for exam in existing_exams:
        supabase.table("generated_exams").delete().eq("id", exam["id"]).execute()

    # 4. Recuperer les etudiants
    student_ids = []
    if session.get("class_id"):
        real_students = list_class_students(session["class_id"])
        student_ids = [{"student_name": s["student_name"], "student_number": s["student_number"]} for s in real_students]
    elif session.get("student_list_id"):
        entries = get_list_entries(session["student_list_id"])
        student_ids = [{"student_name": e["student_name"], "student_number": e["student_number"]} for e in entries]
    else:
        count = session.get("student_count") or 0
        student_ids = [{"student_name": f"Etudiant {i + 1}", "student_number": f"PEAN_{session['id']}_{i + 1:04d}"} for i in range(count)]

    if not student_ids:
        raise HTTPException(
            status_code=400,
            detail="Aucun etudiant dans cette session. Ajoutez des etudiants ou une classe avant de publier.",
        )

    # 5. Structurer le contenu en exercices
    structured = parse_exam_content(content)

    # Si l'enseignant a fourni une configuration manuelle des types, l'utiliser
    if exercises_config:
        try:
            config = json.loads(exercises_config)
            for i, cfg in enumerate(config):
                if i < len(structured):
                    ex_type = cfg.get("type", structured[i].get("exercise_type", "open"))
                    if ex_type not in ("code", "open", "qcm"):
                        ex_type = "open"
                    structured[i]["exercise_type"] = ex_type
                    if ex_type == "code":
                        lang = cfg.get("language")
                        if lang:
                            structured[i]["language"] = lang
                        else:
                            structured[i]["language"] = detect_language(structured[i].get("content", ""))
                    else:
                        structured[i].pop("language", None)
        except (json.JSONDecodeError, IndexError, TypeError) as e:
            logger.warning("exercises_config invalide: %s", e)

    # Si le parseur a détecté une structure cohérente, on l'utilise
    store_content: str = content
    if len(structured) >= 1:
        # Ajouter les champs manquants pour compatibilité avec le frontend étudiant
        for ex in structured:
            ex.setdefault("instructions", "")
            ex.setdefault("points", 10)
            ex.setdefault("difficulty", "medium")
            ex.setdefault("variant_id", 0)
            ex.setdefault("data_overrides", None)
            ex.setdefault("exercise_title", ex.pop("title", f"Question {ex['exercise_id']}"))
        store_content = json.dumps(structured, ensure_ascii=False)

    # 5. Creer une epreuve identique pour chaque etudiant
    total_created = 0
    for student_info in student_ids:
        student_hash = hash_student_identifier(session["id"], student_info["student_number"])
        exam_data = {
            "session_id": session["id"],
            "student_id_hash": student_hash,
            "variant_combo_hash": "shared",
            "sha256_hash": hashlib.sha256(f"{content}:{student_hash}".encode()).hexdigest(),
            "content": store_content,
            "status": "pending",
        }
        created = create_generated_exam(exam_data)
        if created:
            total_created += 1

    # Construire la liste des exercices detectes pour l'affichage enseignant
    detected_exercises = []
    for ex in structured:
        ex_type = ex.get("exercise_type", "open")
        type_label = {"qcm": "QCM", "code": "Code", "open": "Rédaction"}.get(ex_type, ex_type)
        item = {
            "id": ex.get("exercise_id", 0),
            "title": ex.get("exercise_title", "Question"),
            "type": ex_type,
            "type_label": type_label,
            "points": ex.get("points", 10),
        }
        if ex_type == "code":
            item["language"] = ex.get("language", "python")
        detected_exercises.append(item)

    return {
        "generated": total_created,
        "source": source_label,
        "mode": "shared",
        "exercises_created": len(structured),
        "exercises": detected_exercises,
        "message": f"{total_created} epreuves identiques publiees depuis '{source_label}' "
                  f"({len(structured)} question{'' if len(structured) == 1 else 's'} detectee{'' if len(structured) == 1 else 's'})",
    }


@router.get("/{session_id}")
def get_session(
    session_id: int,
    teacher: dict = Depends(get_current_teacher),
):
    """Recuperer une session avec les infos de generation d'epreuves et exercices lies."""
    session = get_session_by_id(session_id)
    if not session or session["teacher_id"] != teacher["id"]:
        raise HTTPException(status_code=404, detail="Session non trouvée")

    result = ExamSessionResponse.model_validate(session).model_dump()
    exams = get_session_exams(session["id"])
    result["exams_generated"] = len(exams)
    result["exams_pending"] = sum(1 for e in exams if e["status"] == "pending")
    result["exams_started"] = sum(1 for e in exams if e["status"] == "started")
    result["exams_submitted"] = sum(1 for e in exams if e["status"] == "submitted")

    # Exercices lies a la session
    linked = get_session_exercises(session["id"])
    result["exercises_count"] = len(linked)
    result["exercises"] = [
        {
            "id": l["id"],
            "exercise_id": l["exercise_id"],
            "sort_order": l["sort_order"],
            "points_override": l.get("points_override"),
            "exercise": l.get("exercises"),
        }
        for l in linked
    ]
    return result


@router.put("/{session_id}", response_model=ExamSessionResponse)
def update_session_route(
    session_id: int,
    data: ExamSessionUpdate,
    teacher: dict = Depends(get_current_teacher),
):
    """Mettre a jour une session d'examen."""
    session = get_session_by_id(session_id)
    if not session or session["teacher_id"] != teacher["id"]:
        raise HTTPException(status_code=404, detail="Session non trouvée")
    if session["status"] != "draft":
        raise HTTPException(
            status_code=400,
            detail="Seules les sessions en brouillon peuvent être modifiées",
        )

    update_data = data.model_dump(exclude_unset=True)
    updated = update_session(session_id, update_data)
    if not updated:
        raise HTTPException(status_code=500, detail="Erreur lors de la mise à jour de la session")
    return ExamSessionResponse.model_validate(updated)


@router.delete("/{session_id}", status_code=204)
def delete_session_route(
    session_id: int,
    teacher: dict = Depends(get_current_teacher),
):
    """Supprimer une session d'examen."""
    session = get_session_by_id(session_id)
    if not session or session["teacher_id"] != teacher["id"]:
        raise HTTPException(status_code=404, detail="Session non trouvée")
    if session["status"] == "active":
        raise HTTPException(
            status_code=400,
            detail="Impossible de supprimer une session active. Terminez-la d'abord.",
        )

    delete_session(session_id)
    return None


# ============================================================
# GESTION DES EXERCICES D'UNE SESSION (session_exercises)
# ============================================================


@router.get("/{session_id}/exercises", response_model=list[SessionExerciseResponse])
def list_session_exercises(
    session_id: int,
    teacher: dict = Depends(get_current_teacher),
):
    """Lister les exercices lies a une session, avec le détail complet."""
    session = get_session_by_id(session_id)
    if not session or session["teacher_id"] != teacher["id"]:
        raise HTTPException(status_code=404, detail="Session non trouvée")

    links = get_session_exercises(session_id)
    return [
        SessionExerciseResponse(
            id=l["id"],
            session_id=l["session_id"],
            exercise_id=l["exercise_id"],
            sort_order=l["sort_order"],
            points_override=l.get("points_override"),
            exercise=l.get("exercises"),
        )
        for l in links
    ]


@router.post("/{session_id}/exercises", status_code=201)
def add_exercise_to_session(
    session_id: int,
    data: SessionExerciseAdd,
    teacher: dict = Depends(get_current_teacher),
):
    """Ajouter un exercice a une session.

    Si l'exercice est deja lie a la session, retourne une erreur 409.
    """
    session = get_session_by_id(session_id)
    if not session or session["teacher_id"] != teacher["id"]:
        raise HTTPException(status_code=404, detail="Session non trouvée")
    if session["status"] != "draft":
        raise HTTPException(status_code=400, detail="Seules les sessions en brouillon peuvent être modifiées")

    # Verifier que l'exercice appartient au professeur
    exercise = get_exercise_by_id(data.exercise_id)
    if not exercise or exercise["teacher_id"] != teacher["id"]:
        raise HTTPException(status_code=404, detail="Exercice non trouvé")

    link = add_session_exercise(
        session_id=session_id,
        exercise_id=data.exercise_id,
        sort_order=data.sort_order,
        points_override=data.points_override,
    )
    if not link:
        raise HTTPException(
            status_code=409,
            detail=f"L'exercice '{exercise['title']}' est déjà dans cette session.",
        )

    return {
        "id": link["id"],
        "exercise_id": link["exercise_id"],
        "exercise_title": exercise["title"],
        "sort_order": link["sort_order"],
        "points_override": link.get("points_override"),
        "message": f"Exercice '{exercise['title']}' ajouté à la session",
    }


@router.delete("/{session_id}/exercises/{exercise_id}", status_code=204)
def remove_exercise_from_session(
    session_id: int,
    exercise_id: int,
    teacher: dict = Depends(get_current_teacher),
):
    """Retirer un exercice d'une session."""
    session = get_session_by_id(session_id)
    if not session or session["teacher_id"] != teacher["id"]:
        raise HTTPException(status_code=404, detail="Session non trouvée")
    if session["status"] != "draft":
        raise HTTPException(status_code=400, detail="Seules les sessions en brouillon peuvent être modifiées")

    removed = remove_session_exercise(session_id, exercise_id)
    if not removed:
        raise HTTPException(status_code=404, detail="Exercice non trouvé dans cette session")
    return None


@router.put("/{session_id}/exercises/reorder")
def reorder_session_exercises(
    session_id: int,
    data: SessionExerciseReorder,
    teacher: dict = Depends(get_current_teacher),
):
    """Reordonner les exercices d'une session.

    Body : { "exercise_ids": [3, 1, 2] }  â† ordre souhaite
    """
    session = get_session_by_id(session_id)
    if not session or session["teacher_id"] != teacher["id"]:
        raise HTTPException(status_code=404, detail="Session non trouvée")
    if session["status"] != "draft":
        raise HTTPException(status_code=400, detail="Seules les sessions en brouillon peuvent être modifiées")

    if not data.exercise_ids:
        raise HTTPException(status_code=400, detail="Liste d'exercices vide")

    success = update_session_exercise_order(session_id, data.exercise_ids)
    if not success:
        raise HTTPException(
            status_code=400,
            detail="Certains exercices ne font pas partie de cette session",
        )
    return {
        "message": "Ordre des exercices mis à jour",
        "exercise_ids": data.exercise_ids,
    }


@router.post("/{session_id}/launch", response_model=ExamSessionResponse)
def launch_session(
    session_id: int,
    teacher: dict = Depends(get_current_teacher),
):
    """Lancer une session d'examen."""
    session = get_session_by_id(session_id)
    if not session or session["teacher_id"] != teacher["id"]:
        raise HTTPException(status_code=404, detail="Session non trouvée")
    if session["status"] != "draft":
        raise HTTPException(
            status_code=400,
            detail="La session n'est pas en mode brouillon",
        )

    updated = update_session(session_id, {"status": "active"})
    if not updated:
        raise HTTPException(status_code=500, detail="Erreur lors du lancement de la session")
    return ExamSessionResponse.model_validate(updated)


@router.post("/{session_id}/complete", response_model=ExamSessionResponse)
def complete_session(
    session_id: int,
    teacher: dict = Depends(get_current_teacher),
):
    """Terminer une session d'examen active."""
    session = get_session_by_id(session_id)
    if not session or session["teacher_id"] != teacher["id"]:
        raise HTTPException(status_code=404, detail="Session non trouvée")
    if session["status"] != "active":
        raise HTTPException(
            status_code=400,
            detail="Seules les sessions actives peuvent être terminées",
        )

    updated = update_session(session_id, {"status": "completed"})
    if not updated:
        raise HTTPException(status_code=500, detail="Erreur lors de la terminaison de la session")
    return ExamSessionResponse.model_validate(updated)


@router.get("/{session_id}/exams/pdf")
def export_exams_pdf(
    session_id: int,
    teacher: dict = Depends(get_current_teacher),
):
    """Export des epreuves generees au format PDF.

    Genere un PDF contenant chaque epreuve etudiante avec les questions,
    choix QCM et enonces. Permet au professeur de verifier le contenu
    avant de partager les epreuves.
    """
    try:
        from fpdf import FPDF
    except ImportError:
        raise HTTPException(status_code=500, detail="fpdf2 n'est pas installe")

    session = get_session_by_id(session_id)
    if not session or session["teacher_id"] != teacher["id"]:
        raise HTTPException(status_code=404, detail="Session non trouvee")

    # 1. Recuperer les epreuves generees
    exams = get_session_exams(session_id)
    if not exams:
        raise HTTPException(status_code=404, detail="Aucune epreuve generee pour cette session")

    # 2. Recuperer la liste des etudiants pour associer noms aux epreuves
    student_map: dict[str, dict] = {}
    if session.get("class_id"):
        students = list_class_students(session["class_id"])
        for s in students:
            h = hash_student_identifier(session_id, s["student_number"])
            student_map[h] = s
    elif session.get("student_list_id"):
        entries = get_list_entries(session["student_list_id"])
        for e in entries:
            h = hash_student_identifier(session_id, e["student_number"])
            student_map[h] = e

    class ExamsPDF(FPDF):
        def header(self):
            if self.page_no() > 1:
                self.set_font("Helvetica", "I", 8)
                self.set_text_color(128)
                self.cell(0, 5, f"PEAN - {session['title']}", align="C", new_x="LMARGIN", new_y="NEXT")
                self.ln(2)

        def footer(self):
            self.set_y(-15)
            self.set_font("Helvetica", "I", 8)
            self.set_text_color(128)
            self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", align="C")

    pdf = ExamsPDF()
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=20)

    for idx, exam in enumerate(exams):
        # Titre de l'epreuve
        pdf.add_page()

        # En-tete de l'epreuve
        pdf.set_font("Helvetica", "B", 16)
        pdf.set_text_color(15, 23, 42)
        pdf.cell(0, 10, f"Epreuve {idx + 1}", align="C", new_x="LMARGIN", new_y="NEXT")

        # Information etudiant
        student_hash = exam.get("student_id_hash", "")
        student_info = student_map.get(student_hash)
        if student_info:
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(100, 116, 139)
            pdf.cell(0, 6, f"Etudiant: {student_info.get('student_name', 'N/A')}", new_x="LMARGIN", new_y="NEXT")
            pdf.cell(0, 6, f"Matricule: {student_info.get('student_number', 'N/A')}", new_x="LMARGIN", new_y="NEXT")

        pdf.set_draw_color(37, 99, 235)
        pdf.set_line_width(0.5)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(5)

        # Parser le contenu JSON
        try:
            exercises = json.loads(exam.get("content", "[]"))
        except (json.JSONDecodeError, TypeError):
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(220, 38, 38)
            pdf.cell(0, 10, "Erreur: contenu de l'epreuve invalide", new_x="LMARGIN", new_y="NEXT")
            continue

        if not exercises:
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(100, 116, 139)
            pdf.cell(0, 10, "Cette epreuve est vide.", new_x="LMARGIN", new_y="NEXT")
            continue

        for q_idx, q in enumerate(exercises):
            # Verifier si on doit ajouter une page
            if pdf.get_y() > 240:
                pdf.add_page()
                pdf.set_font("Helvetica", "B", 16)
                pdf.set_text_color(15, 23, 42)
                pdf.cell(0, 10, f"Epreuve {idx + 1} (suite)", align="C", new_x="LMARGIN", new_y="NEXT")
                pdf.ln(3)

            # Question numero
            pdf.set_font("Helvetica", "B", 12)
            pdf.set_text_color(15, 23, 42)
            points = q.get("points", "N/A")
            ex_type = {
                "qcm": "QCM",
                "open": "Redaction",
                "code": "Code",
            }.get(q.get("exercise_type", ""), "")
            pdf.cell(0, 8, f"Question {q_idx + 1}  |  {points} pts  |  {ex_type}", new_x="LMARGIN", new_y="NEXT")

            # Instructions / enonce
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(55, 65, 81)
            content = q.get("content") or q.get("instructions") or ""
            # Decouper le texte en lignes
            pdf.multi_cell(0, 5, content)
            pdf.ln(2)

            # Choix QCM
            data_overrides = q.get("data_overrides")
            if data_overrides and isinstance(data_overrides, dict):
                choices = data_overrides.get("choices", [])
                if choices:
                    pdf.set_font("Helvetica", "B", 10)
                    pdf.set_text_color(37, 99, 235)
                    pdf.cell(0, 6, "Choix:", new_x="LMARGIN", new_y="NEXT")
                    pdf.set_font("Helvetica", "", 10)
                    pdf.set_text_color(55, 65, 81)
                    for choice in choices:
                        pdf.cell(5)
                        pdf.multi_cell(0, 5, choice)
                    pdf.ln(2)

                # Cas de test pour le code
                test_cases = data_overrides.get("test_cases", [])
                if test_cases:
                    pdf.set_font("Helvetica", "B", 10)
                    pdf.set_text_color(37, 99, 235)
                    pdf.cell(0, 6, "Cas de test:", new_x="LMARGIN", new_y="NEXT")
                    pdf.set_font("Helvetica", "", 9)
                    pdf.set_text_color(55, 65, 81)
                    for tc in test_cases:
                        inp = tc.get("input", "")
                        expected = tc.get("expected_output", "")
                        pdf.cell(5)
                        pdf.multi_cell(0, 4, f"Entree: {inp}")
                        pdf.cell(5)
                        pdf.multi_cell(0, 4, f"Attendu: {expected}")
                    pdf.ln(2)

            pdf.ln(4)
            pdf.set_draw_color(226, 232, 240)
            pdf.set_line_width(0.2)
            pdf.line(15, pdf.get_y(), 200, pdf.get_y())
            pdf.ln(4)

        # Espacement entre epreuves
        pdf.ln(5)

    output = io.BytesIO()
    pdf.output(output)
    output.seek(0)

    return StreamingResponse(
        output,
        media_type="application/pdf",
        headers={
            "Content-Disposition": f"attachment; filename=session_{session_id}_epreuves.pdf"
        },
    )

